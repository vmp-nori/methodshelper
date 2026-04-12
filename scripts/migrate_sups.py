"""
migrate_sups.py — Pulls MM SUP .docx files from the nori.study Supabase project,
parses the exercise/question lists, and uploads them to the SUPsmasher sup_data table.

Requirements:
    pip install supabase python-docx

Usage:
    set SOURCE_SERVICE_KEY=<nori.study service-role key>
    set DEST_SERVICE_KEY=<SUPsmasher service-role key>
    python scripts/migrate_sups.py
"""

import os
import re
import json

# ── Config ───────────────────────────────────────────────────────────────────
SOURCE_URL    = "https://rrpfggsrqlhyezcuyqmb.supabase.co"   # nori.study
SOURCE_KEY    = os.environ.get("SOURCE_SERVICE_KEY", "")

DEST_URL      = "https://bodwvpqtfhqmpzialpig.supabase.co"   # SUPsmasher
DEST_KEY      = os.environ.get("DEST_SERVICE_KEY", "")

SOURCE_BUCKET = "methods"
DEST_SUBJECT  = "MM12"   # subject code in SUPsmasher to match against
# ─────────────────────────────────────────────────────────────────────────────


def get_clients():
    from supabase import create_client
    if not SOURCE_KEY:
        raise SystemExit("Set SOURCE_SERVICE_KEY env var (nori.study service-role key)")
    if not DEST_KEY:
        raise SystemExit("Set DEST_SERVICE_KEY env var (SUPsmasher service-role key)")
    return create_client(SOURCE_URL, SOURCE_KEY), create_client(DEST_URL, DEST_KEY)


def get_subject_id(client, url: str, code: str) -> str:
    result = client.table("subjects").select("id").eq("code", code).single().execute()
    if not result.data:
        raise SystemExit(f"Subject '{code}' not found in {url}")
    return result.data["id"]


def fetch_mm_sups(source) -> list[dict]:
    """Fetch all MM unit_outline resources from nori.study."""
    # Get MM subject id in nori.study (code is "MM" there, not "MM12")
    subj_result = source.table("subjects").select("id, code").execute()
    mm_ids = [r["id"] for r in subj_result.data if r["code"].startswith("MM")]
    if not mm_ids:
        raise SystemExit("No MM subjects found in nori.study")

    result = (
        source.table("resources")
        .select("id, name, file_path, topic")
        .in_("subject_id", mm_ids)
        .eq("resource_type", "unit_outline")
        .order("topic")
        .execute()
    )
    return result.data


def download_docx(source, file_path: str) -> bytes:
    return source.storage.from_(SOURCE_BUCKET).download(file_path)


def parse_exercise_content(content: str) -> list[int]:
    """
    Parse question numbers from the content following an exercise code.

    Handles all MM SUP formats observed across T00-T10:
      Q-prefix:        "Q3, Q4a, c, Q5"       -> [3, 4, 5]
      Q then bare:     "Q1, 3, 4, 5, 7"        -> [1, 3, 4, 5, 7]
      Q-range:         "Q1 - Q6" or "Q1-2"     -> [1,2,3,4,5,6] / [1,2]
      Bare numbers:    "1, 2, 3, 5, 7"         -> [1, 2, 3, 5, 7]
      Number ranges:   "1-8, 10-13"            -> [1..8, 10..13]
      ESO:             "10-13 ESO"             -> [10, 12]
      ALL:             "ALL"                   -> [] (whole exercise, skipped)

    Sub-part letters like "Q4a, c, e" are ignored -- only integers kept.
    """
    nums: set[int] = set()

    # Strip parenthetical notes: "(Every Second One)", "(use tree diagram)", etc.
    working = re.sub(r'\([^)]*\)', '', content)
    # Strip whole-exercise ALL
    working = re.sub(r'\bALL\b', '', working, flags=re.IGNORECASE)

    # Pass 1: Q-ranges  "Q1 - Q6"  or  "Q1-2"
    for m in re.finditer(r'Q\s*(\d+)\s*[-\u2013\u2014]+\s*Q?\s*(\d+)', working, re.IGNORECASE):
        lo, hi = int(m.group(1)), int(m.group(2))
        nums.update(range(lo, hi + 1))
    working = re.sub(r'Q\s*\d+\s*[-\u2013\u2014]+\s*Q?\s*\d+', '', working, flags=re.IGNORECASE)

    # Pass 2: explicit Q-prefixed numbers
    for m in re.finditer(r'Q\s*(\d+)', working, re.IGNORECASE):
        nums.add(int(m.group(1)))
    working = re.sub(r'Q\s*\d+', '', working, flags=re.IGNORECASE)

    # Pass 3: bare numbers and ranges per comma/semicolon token
    for token in re.split(r'[,;]', working):
        token = token.strip()
        if not token:
            continue

        eso = bool(re.search(r'\bESO\b', token, re.IGNORECASE))
        token = re.sub(r'\bESO\b', '', token, flags=re.IGNORECASE).strip()

        # Number range like "1-8" or "10-13"
        rm = re.search(r'(\d+)\s*[-\u2013\u2014]\s*(\d+)', token)
        if rm:
            lo, hi = int(rm.group(1)), int(rm.group(2))
            step = 2 if eso else 1
            nums.update(range(lo, hi + 1, step))
            continue

        # Single number, possibly with trailing sub-part letters like "9a" -- cap at 99
        nm = re.search(r'(\d+)', token)
        if nm:
            n = int(nm.group(1))
            if 1 <= n <= 99:
                nums.add(n)

    return sorted(nums)


def parse_docx(docx_bytes: bytes) -> list[dict]:
    """
    Parse a SUP .docx and extract exercise + question-number lists.

    SUPs use a table with a "Test Yourself" / "Text Book Work" column.
    Entries follow one of two styles:
        Q-prefix:    Ex 2B - Q3, Q4a, Q5, Q6
        Bare-number: Ex 13A - 1-8, 9a, c, f, h, 10-13 ESO

    Returns: [{"exercise": "2B", "questions": [3, 4, 5, 6]}, ...]
    """
    from docx import Document
    import io

    doc = Document(io.BytesIO(docx_bytes))

    ex_pattern = re.compile(r'Ex\s+(\d+[A-Za-z]+)', re.IGNORECASE)

    exercises_map: dict[str, set] = {}

    def process_line(line: str):
        # Split on Ex markers; result is [before, code, after, code, after, ...]
        segments = ex_pattern.split(line)
        i = 1
        while i < len(segments) - 1:
            ex_code = segments[i].upper()
            after   = segments[i + 1]
            q_nums  = parse_exercise_content(after)
            if q_nums:
                exercises_map.setdefault(ex_code, set()).update(q_nums)
            i += 2

    def process_text(text: str):
        for line in text.splitlines():
            line = line.strip()
            if line:
                process_line(line)

    for para in doc.paragraphs:
        process_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text(cell.text)

    return [
        {"exercise": ex, "questions": sorted(qs)}
        for ex, qs in sorted(exercises_map.items())
        if qs
    ]


def upload_sup_data(dest, subject_id: str, topic_code: str, topic_name: str, exercises: list[dict]):
    dest.table("sup_data").upsert({
        "subject_id": subject_id,
        "topic_code": topic_code,
        "topic_name": topic_name,
        "exercises":  exercises,
    }, on_conflict="subject_id,topic_code").execute()
    print(f"  [ok] Uploaded {len(exercises)} exercises for {topic_code}")


def main():
    source, dest = get_clients()

    dest_subject_id = get_subject_id(dest, DEST_URL, DEST_SUBJECT)
    print(f"SUPsmasher MM12 subject ID: {dest_subject_id}\n")

    records = fetch_mm_sups(source)
    print(f"Found {len(records)} MM SUP files in nori.study\n")

    for rec in records:
        topic_field = rec.get("topic", "")
        code_match  = re.match(r'(T\d+)', topic_field)
        topic_code  = code_match.group(1) if code_match else topic_field
        topic_name  = re.sub(r'^T\d+\s*[-\u2013]\s*', '', topic_field).strip()

        print(f"Processing {topic_code} — {topic_name}")
        print(f"  File: {rec['file_path']}")

        try:
            docx_bytes = download_docx(source, rec["file_path"])
            exercises  = parse_docx(docx_bytes)

            if not exercises:
                print(f"  [!] No exercises parsed -- check the docx format manually")
                print()
                continue

            for ex in exercises:
                print(f"    Exercise {ex['exercise']}: {len(ex['questions'])} questions -> {ex['questions']}")

            upload_sup_data(dest, dest_subject_id, topic_code, topic_name, exercises)

        except Exception as e:
            print(f"  [x] Error: {e}")

        print()

    print("Migration complete.")


if __name__ == "__main__":
    main()
