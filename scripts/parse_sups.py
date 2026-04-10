"""
parse_sups.py — Downloads SUP .docx files from Supabase Storage,
parses the exercise/question lists, and uploads structured data
to the sup_data table.

Requirements:
    pip install supabase python-docx

Usage:
    python parse_sups.py

Set your Supabase service-role key in the SUPABASE_SERVICE_KEY env var,
or paste it directly into the config section below.
"""

import os
import re
import json
import tempfile
from pathlib import Path

# ── Config ──────────────────────────────────────────────────────────────────
SUPABASE_URL     = "https://bodwvpqtfhqmpzialpig.supabase.co"
SUPABASE_KEY     = os.environ.get("SUPABASE_SERVICE_KEY", "")   # service-role key
STORAGE_BUCKET   = "methods"
# ────────────────────────────────────────────────────────────────────────────


def get_supabase():
    from supabase import create_client
    if not SUPABASE_KEY:
        raise SystemExit("Set SUPABASE_SERVICE_KEY env var (service-role key from Supabase dashboard)")
    return create_client(SUPABASE_URL, SUPABASE_KEY)


def get_subject_id(client, code: str) -> str:
    result = client.table("subjects").select("id").eq("code", code).single().execute()
    if not result.data:
        raise SystemExit(f"Subject '{code}' not found in database.")
    return result.data["id"]


def download_docx(client, file_path: str) -> bytes:
    """Download a file from Supabase Storage and return its bytes."""
    response = client.storage.from_(STORAGE_BUCKET).download(file_path)
    return response


def parse_docx(docx_bytes: bytes) -> list[dict]:
    """
    Parse a SUP .docx file and extract exercise + question-number lists.

    Expected document structure (may vary):
        Exercise 1A   Questions: 1, 3, 5, 7, 9
        Exercise 1B   Questions: 2, 4, 6, 8

    Returns:
        [{"exercise": "1A", "questions": [1, 3, 5, 7, 9]}, ...]
    """
    from docx import Document
    import io

    doc = Document(io.BytesIO(docx_bytes))
    exercises = []

    # Regex patterns
    # Matches "Exercise 3A" or "Ex 3A" or just "3A" as a label
    ex_pattern  = re.compile(r'Exercise\s+(\d+[A-Za-z]+)', re.IGNORECASE)
    # Matches a list of numbers/ranges like "1, 3, 5-9, 11"
    num_pattern = re.compile(r'(?:Questions?[:\s]*)?([\d,\s\-–]+)', re.IGNORECASE)

    current_exercise = None
    current_questions = []

    def flush():
        if current_exercise and current_questions:
            exercises.append({
                "exercise": current_exercise,
                "questions": sorted(set(current_questions))
            })

    def expand_range(token: str) -> list[int]:
        """Expand "5-9" or "5–9" into [5, 6, 7, 8, 9]."""
        token = token.strip().replace('–', '-')
        if '-' in token:
            parts = token.split('-')
            try:
                lo, hi = int(parts[0].strip()), int(parts[1].strip())
                return list(range(lo, hi + 1))
            except ValueError:
                return []
        try:
            return [int(token)]
        except ValueError:
            return []

    def extract_numbers(text: str) -> list[int]:
        """Pull all question numbers out of a text fragment."""
        nums = []
        for token in re.split(r'[,\s]+', text):
            nums.extend(expand_range(token))
        return [n for n in nums if n > 0]

    # Walk every paragraph and every table cell
    def process_text(text: str):
        nonlocal current_exercise, current_questions
        text = text.strip()
        if not text:
            return

        ex_match = ex_pattern.search(text)
        if ex_match:
            flush()
            current_exercise = ex_match.group(1).upper()
            current_questions = []
            # Numbers might be on the same line after the exercise label
            rest = text[ex_match.end():]
            nums = extract_numbers(rest)
            current_questions.extend(nums)
        elif current_exercise:
            # If we already have an exercise, try to pick up question numbers
            nums = extract_numbers(text)
            if nums:
                current_questions.extend(nums)

    for para in doc.paragraphs:
        process_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_text(cell.text)

    flush()
    return exercises


def upload_sup_data(client, subject_id: str, topic_code: str, topic_name: str, exercises: list[dict]):
    """Upsert SUP data into the sup_data table."""
    client.table("sup_data").upsert({
        "subject_id":  subject_id,
        "topic_code":  topic_code,
        "topic_name":  topic_name,
        "exercises":   json.dumps(exercises),
    }, on_conflict="subject_id,topic_code").execute()
    print(f"  ✓ Uploaded {len(exercises)} exercises for {topic_code}")


def main():
    client     = get_supabase()
    subject_id = get_subject_id(client, "MM12")

    # Fetch all MM SUP resource records
    result = (
        client.table("resources")
        .select("id, name, file_path, topic")
        .eq("subject_id", subject_id)
        .eq("resource_type", "unit_outline")
        .order("topic")
        .execute()
    )
    records = result.data
    print(f"Found {len(records)} SUP files to process\n")

    for rec in records:
        topic_field = rec.get("topic", "")
        # Extract topic code like "T01" from "T01 - Coordinate Geometry"
        code_match = re.match(r'(T\d+)', topic_field)
        topic_code = code_match.group(1) if code_match else topic_field
        topic_name = re.sub(r'^T\d+\s*[-–]\s*', '', topic_field).strip()

        print(f"Processing {topic_code} — {topic_name}")
        print(f"  File: {rec['file_path']}")

        try:
            docx_bytes = download_docx(client, rec["file_path"])
            exercises  = parse_docx(docx_bytes)

            if not exercises:
                print(f"  ⚠  No exercises found — check the docx format manually")
                continue

            for ex in exercises:
                print(f"    Exercise {ex['exercise']}: {len(ex['questions'])} questions → {ex['questions']}")

            upload_sup_data(client, subject_id, topic_code, topic_name, exercises)

        except Exception as e:
            print(f"  ✗ Error: {e}")

        print()

    print("Done.")


if __name__ == "__main__":
    main()
